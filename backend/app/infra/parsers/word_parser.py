"""Word parser — extracts text from .docx files using python-docx."""

import hashlib
from pathlib import Path

from app.core.exceptions import ParseError
from app.core.logging import get_logger

logger = get_logger(__name__)

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore[assignment]


class WordParser:
    """Extract text content from Word (.docx) files."""

    supported_extensions = {".docx"}

    async def parse(self, file_path: str | Path) -> str:
        """Read a .docx file and return its full text content."""
        if Document is None:
            raise ParseError(
                "python-docx is not installed. Run: pip install python-docx"
            )

        path = Path(file_path)
        if not path.exists():
            raise ParseError(f"File not found: {path}")

        try:
            doc = Document(str(path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            full_text = "\n\n".join(paragraphs)
            logger.info(f"Parsed Word: {path.name}, {len(paragraphs)} paragraphs, {len(full_text)} chars")
            return full_text

        except Exception as e:
            raise ParseError(f"Failed to parse Word document '{path.name}': {e}")

    def compute_hash(self, file_path: str | Path) -> str:
        """SHA-256 hash of the file contents."""
        path = Path(file_path)
        sha = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha.update(chunk)
        return sha.hexdigest()
